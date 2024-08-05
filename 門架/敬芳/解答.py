import docx

def extract_questions_and_answers(doc_path):
    doc = docx.Document(doc_path)
    questions = []

    tables = doc.tables
    for table in tables:
        for idx, row in enumerate(table.rows):
            # 假設第一行是標題，跳過
            if idx == 0:
                continue
            
            # 假設固定有四個欄位，分別是序號、課程名稱、題目、答案
            serial_number = row.cells[0].text.strip()
            course_name = row.cells[1].text.strip()
            question = row.cells[2].text.strip()
            answer = row.cells[3].text.strip()
            
            # 提取問題文本（假設問題以 (A) 結束）
            question_text = question.split('(A)')[0].strip()
            
            # 找到答案對應的選項
            options = question.split('(')[1:]
            answer_option = ''
            for option in options:
                if option.startswith(answer):
                    answer_option = f"({answer})" + option.split(')')[1]
                    break
            
            # 合併問題和答案，包含原始序號
            combined = f"{serial_number}. {question_text} {answer_option}"
            questions.append(combined)
    
    return questions

def save_to_docx(questions, output_path):
    doc = docx.Document()
    
    for question in questions:
        doc.add_paragraph(question)
    
    doc.save(output_path)

# 使用範例
doc_path = "/workspaces/-/ACFrOgBNXKn6OajETtYrWKr5kwDF9CHLna_KAb-Azb-t-642g6JKkRRZ8GO34D44xB3YJvKQV4B1tTTihqKHIWx08rZxaMPbHQkKi6WvT5owIDByUHhbKuK.docx"
output_path = "/workspaces/-/修改-答案.docx"

questions_and_answers = extract_questions_and_answers(doc_path)
save_to_docx(questions_and_answers, output_path)

print('完成！')
